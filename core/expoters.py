from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
import json
from io import BytesIO
from django.http import HttpResponse

class QuestionExporter:
    def __init__(self, questions):
        self.questions = questions
    
    def export_pdf(self):
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center
        )
        story.append(Paragraph("Generated Questions", title_style))
        story.append(Spacer(1, 20))
        
        # Group questions by Bloom level
        bloom_groups = {}
        for question in self.questions:
            bloom_level = question.get_bloom_level_display()
            if bloom_level not in bloom_groups:
                bloom_groups[bloom_level] = []
            bloom_groups[bloom_level].append(question)
        
        for bloom_level, questions in bloom_groups.items():
            # Bloom level header
            story.append(Paragraph(f"{bloom_level} Questions", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            for i, question in enumerate(questions, 1):
                # Question number and text
                story.append(Paragraph(f"Q{i}. {question.question_text}", styles['Normal']))
                
                # Options for multiple choice
                if question.options:
                    for j, option in enumerate(question.options):
                        story.append(Paragraph(f"   {chr(65+j)}. {option}", styles['Normal']))
                
                # Answer section
                if question.answer:
                    story.append(Paragraph(f"<b>Answer:</b> {question.answer}", styles['Normal']))
                
                if question.explanation:
                    story.append(Paragraph(f"<b>Explanation:</b> {question.explanation}", styles['Normal']))
                
                story.append(Spacer(1, 12))
            
            story.append(Spacer(1, 20))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def export_docx(self):
        doc = Document()
        
        # Add title
        title = doc.add_heading('Generated Questions', 0)
        title.alignment = 1  # Center
        
        # Group questions by Bloom level
        bloom_groups = {}
        for question in self.questions:
            bloom_level = question.get_bloom_level_display()
            if bloom_level not in bloom_groups:
                bloom_groups[bloom_level] = []
            bloom_groups[bloom_level].append(question)
        
        for bloom_level, questions in bloom_groups.items():
            # Add Bloom level header
            doc.add_heading(f"{bloom_level} Questions", level=1)
            
            for i, question in enumerate(questions, 1):
                # Question
                q_para = doc.add_paragraph()
                q_para.add_run(f"Q{i}. ").bold = True
                q_para.add_run(question.question_text)
                
                # Options for multiple choice
                if question.options:
                    for j, option in enumerate(question.options):
                        option_para = doc.add_paragraph(f"   {chr(65+j)}. {option}")
                        option_para.style = 'List Bullet'
                
                # Answer
                if question.answer:
                    answer_para = doc.add_paragraph()
                    answer_para.add_run("Answer: ").bold = True
                    answer_para.add_run(question.answer)
                
                # Explanation
                if question.explanation:
                    exp_para = doc.add_paragraph()
                    exp_para.add_run("Explanation: ").bold = True
                    exp_para.add_run(question.explanation)
                
                doc.add_paragraph()  # Add space
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def export_json(self):
        data = {
            'questions': [
                {
                    'id': str(question.id),
                    'question_text': question.question_text,
                    'question_type': question.question_type,
                    'bloom_level': question.bloom_level,
                    'difficulty': question.difficulty,
                    'answer': question.answer,
                    'options': question.options,
                    'explanation': question.explanation,
                    'topic': question.topic
                }
                for question in self.questions
            ],
            'metadata': {
                'total_questions': len(self.questions),
                'bloom_levels': list(set(q.bloom_level for q in self.questions)),
                'question_types': list(set(q.question_type for q in self.questions))
            }
        }
        return json.dumps(data, indent=2)
