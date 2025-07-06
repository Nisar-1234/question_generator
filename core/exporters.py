# core/exporters.py
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import json
from io import BytesIO
from django.http import HttpResponse
from django.utils import timezone
import csv
from docx.shared import RGBColor


class QuestionExporter:
    def __init__(self, questions, file_obj=None):
        self.questions = questions
        self.file_obj = file_obj
        self.total_questions = questions.count()
    
    def export_pdf(self):
        """Export questions as professional PDF"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=A4,
            rightMargin=72, 
            leftMargin=72, 
            topMargin=72, 
            bottomMargin=50
        )
        
        # Container for PDF elements
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center
            textColor=colors.darkblue,
            fontName='Helvetica-Bold'
        )
        
        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=20,
            alignment=1,
            textColor=colors.grey
        )
        
        bloom_heading_style = ParagraphStyle(
            'BloomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            spaceBefore=20,
            textColor=colors.darkblue,
            borderWidth=1,
            borderColor=colors.darkblue,
            borderPadding=5,
            backColor=colors.lightblue,
            fontName='Helvetica-Bold'
        )
        
        question_style = ParagraphStyle(
            'Question',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=8,
            leftIndent=0,
            fontName='Helvetica-Bold'
        )
        
        answer_style = ParagraphStyle(
            'Answer',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=8,
            leftIndent=20,
            textColor=colors.darkgreen,
            fontName='Helvetica'
        )
        
        explanation_style = ParagraphStyle(
            'Explanation',
            parent=styles['Normal'],
            fontSize=9,
            spaceAfter=15,
            leftIndent=20,
            textColor=colors.grey,
            fontName='Helvetica-Oblique'
        )
        
        # Title page
        elements.append(Paragraph("Generated Questions & Answers", title_style))
        
        if self.file_obj:
            elements.append(Paragraph(f"Document: {self.file_obj.filename}", subtitle_style))
            elements.append(Paragraph(f"Generated on: {timezone.now().strftime('%B %d, %Y at %I:%M %p')}", subtitle_style))
        
        elements.append(Paragraph(f"Total Questions: {self.total_questions}", subtitle_style))
        elements.append(Spacer(1, 30))
        
        # Table of contents
        elements.append(Paragraph("Table of Contents", styles['Heading2']))
        bloom_levels = self.questions.values_list('bloom_level', flat=True).distinct()
        for bloom in bloom_levels:
            count = self.questions.filter(bloom_level=bloom).count()
            elements.append(Paragraph(f"• {bloom.replace('_', ' ').title()} Questions ({count})", styles['Normal']))
        
        elements.append(PageBreak())
        
        # Group questions by Bloom level
        current_bloom = None
        question_num = 1
        
        for question in self.questions.order_by('bloom_level', 'question_type', 'created_at'):
            # Add Bloom level header
            if current_bloom != question.bloom_level:
                current_bloom = question.bloom_level
                bloom_display = question.get_bloom_level_display()
                elements.append(Paragraph(f"{bloom_display} Questions", bloom_heading_style))
                elements.append(Spacer(1, 15))
            
            # Question number and text
            question_text = f"Question {question_num}: {question.question_text}"
            elements.append(Paragraph(question_text, question_style))
            
            # Question metadata
            metadata = f"Type: {question.get_question_type_display()} | Difficulty: {question.get_difficulty_display()}"
            elements.append(Paragraph(metadata, styles['Normal']))
            elements.append(Spacer(1, 8))
            
            # Options for multiple choice
            if question.options and len(question.options) > 0:
                elements.append(Paragraph("Options:", styles['Normal']))
                for i, option in enumerate(question.options):
                    option_text = f"   {chr(65+i)}. {option}"
                    elements.append(Paragraph(option_text, styles['Normal']))
                elements.append(Spacer(1, 8))
            
            # Answer
            if question.answer:
                answer_text = f"<b>Answer:</b> {question.answer}"
                elements.append(Paragraph(answer_text, answer_style))
            
            # Explanation
            if question.explanation:
                explanation_text = f"<b>Explanation:</b> {question.explanation}"
                elements.append(Paragraph(explanation_text, explanation_style))
            
            elements.append(Spacer(1, 20))
            question_num += 1
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        return buffer
    
    def export_docx(self):
        """Export questions as professional Word document"""
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title
        title = doc.add_heading('Generated Questions & Answers', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Document information
        if self.file_obj:
            info_para = doc.add_paragraph()
            info_para.add_run(f"Document: ").bold = True
            info_para.add_run(f"{self.file_obj.filename}")
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info_para2 = doc.add_paragraph()
        info_para2.add_run(f"Generated on: ").bold = True
        info_para2.add_run(f"{timezone.now().strftime('%B %d, %Y at %I:%M %p')}")
        info_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        total_para = doc.add_paragraph()
        total_para.add_run(f"Total Questions: ").bold = True
        total_para.add_run(f"{self.total_questions}")
        total_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Empty line
        
        # Table of contents
        toc_heading = doc.add_heading('Table of Contents', level=1)
        bloom_levels = self.questions.values_list('bloom_level', flat=True).distinct()
        for bloom in bloom_levels:
            count = self.questions.filter(bloom_level=bloom).count()
            toc_para = doc.add_paragraph()
            toc_para.add_run(f"• {bloom.replace('_', ' ').title()} Questions ({count})")
        
        doc.add_page_break()
        
        # Group questions by Bloom level
        current_bloom = None
        question_num = 1
        
        for question in self.questions.order_by('bloom_level', 'question_type', 'created_at'):
            # Add Bloom level header
            if current_bloom != question.bloom_level:
                current_bloom = question.bloom_level
                bloom_display = question.get_bloom_level_display()
                bloom_heading = doc.add_heading(f"{bloom_display} Questions", level=1)
                bloom_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
            
            # Question
            q_para = doc.add_paragraph()
            q_para.add_run(f"Question {question_num}: ").bold = True
            q_para.add_run(question.question_text)
            
            # Question metadata
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"Type: {question.get_question_type_display()} | ")
            meta_para.add_run(f"Difficulty: {question.get_difficulty_display()}")
            meta_para.runs[0].italic = True
            
            # Options for multiple choice
            if question.options and len(question.options) > 0:
                options_para = doc.add_paragraph()
                options_para.add_run("Options:").bold = True
                
                for i, option in enumerate(question.options):
                    option_para = doc.add_paragraph(f"   {chr(65+i)}. {option}")
                    option_para.style = 'List Bullet'
            
            # Answer
            if question.answer:
                answer_para = doc.add_paragraph()
                answer_para.add_run("Answer: ").bold = True
                answer_run = answer_para.add_run(question.answer)
                answer_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            
            # Explanation
            if question.explanation:
                exp_para = doc.add_paragraph()
                exp_para.add_run("Explanation: ").bold = True
                exp_run = exp_para.add_run(question.explanation)
                exp_run.italic = True
                exp_run.font.color.rgb = RGBColor(102, 102, 102)  # Gray
            
            doc.add_paragraph()  # Empty line between questions
            question_num += 1
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def export_txt(self):
        """Export questions as formatted text file"""
        content = []
        content.append("=" * 80)
        content.append("GENERATED QUESTIONS & ANSWERS")
        content.append("=" * 80)
        content.append("")
        
        if self.file_obj:
            content.append(f"Document: {self.file_obj.filename}")
        content.append(f"Generated on: {timezone.now().strftime('%B %d, %Y at %I:%M %p')}")
        content.append(f"Total Questions: {self.total_questions}")
        content.append("")
        
        # Table of contents
        content.append("TABLE OF CONTENTS")
        content.append("-" * 40)
        bloom_levels = self.questions.values_list('bloom_level', flat=True).distinct()
        for bloom in bloom_levels:
            count = self.questions.filter(bloom_level=bloom).count()
            content.append(f"• {bloom.replace('_', ' ').title()} Questions ({count})")
        content.append("")
        content.append("=" * 80)
        content.append("")
        
        # Questions grouped by Bloom level
        current_bloom = None
        question_num = 1
        
        for question in self.questions.order_by('bloom_level', 'question_type', 'created_at'):
            # Bloom level header
            if current_bloom != question.bloom_level:
                current_bloom = question.bloom_level
                bloom_display = question.get_bloom_level_display()
                content.append("")
                content.append(f"{bloom_display.upper()} QUESTIONS")
                content.append("=" * len(f"{bloom_display.upper()} QUESTIONS"))
                content.append("")
            
            # Question
            content.append(f"Question {question_num}:")
            content.append(f"{question.question_text}")
            content.append("")
            
            # Metadata
            content.append(f"Type: {question.get_question_type_display()}")
            content.append(f"Difficulty: {question.get_difficulty_display()}")
            content.append(f"Bloom Level: {question.get_bloom_level_display()}")
            content.append("")
            
            # Options
            if question.options and len(question.options) > 0:
                content.append("Options:")
                for i, option in enumerate(question.options):
                    content.append(f"   {chr(65+i)}. {option}")
                content.append("")
            
            # Answer
            if question.answer:
                content.append(f"ANSWER: {question.answer}")
                content.append("")
            
            # Explanation
            if question.explanation:
                content.append(f"EXPLANATION: {question.explanation}")
                content.append("")
            
            content.append("-" * 60)
            content.append("")
            question_num += 1
        
        return "\n".join(content)
    
    def export_json(self):
        """Export questions as structured JSON"""
        data = {
            "export_info": {
                "generated_on": timezone.now().isoformat(),
                "total_questions": self.total_questions,
                "document_filename": self.file_obj.filename if self.file_obj else "Multiple files",
                "export_format": "complete_questions_and_answers"
            },
            "summary": {
                "questions_by_bloom_level": {},
                "questions_by_type": {},
                "questions_by_difficulty": {}
            },
            "questions": []
        }
        
        # Generate summary statistics
        for bloom in self.questions.values_list('bloom_level', flat=True).distinct():
            data["summary"]["questions_by_bloom_level"][bloom] = self.questions.filter(bloom_level=bloom).count()
        
        for qtype in self.questions.values_list('question_type', flat=True).distinct():
            data["summary"]["questions_by_type"][qtype] = self.questions.filter(question_type=qtype).count()
        
        for difficulty in self.questions.values_list('difficulty', flat=True).distinct():
            data["summary"]["questions_by_difficulty"][difficulty] = self.questions.filter(difficulty=difficulty).count()
        
        # Export all questions
        for i, question in enumerate(self.questions.order_by('bloom_level', 'question_type', 'created_at'), 1):
            question_data = {
                "question_number": i,
                "id": str(question.id),
                "question_text": question.question_text,
                "question_type": question.question_type,
                "question_type_display": question.get_question_type_display(),
                "bloom_level": question.bloom_level,
                "bloom_level_display": question.get_bloom_level_display(),
                "difficulty": question.difficulty,
                "difficulty_display": question.get_difficulty_display(),
                "answer": question.answer,
                "options": question.options if question.options else [],
                "explanation": question.explanation,
                "topic": question.topic,
                "created_at": question.created_at.isoformat(),
                "file_info": {
                    "filename": question.file.filename,
                    "file_id": str(question.file.id)
                }
            }
            data["questions"].append(question_data)
        
        return json.dumps(data, indent=2, ensure_ascii=False)
    
    def export_csv(self):
        """Export questions as CSV for Excel"""
        buffer = BytesIO()
        # Use StringIO for CSV in memory
        import io
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Headers
        writer.writerow([
            'Question Number',
            'Question Text',
            'Question Type',
            'Bloom Level',
            'Difficulty',
            'Answer',
            'Options (separated by |)',
            'Explanation',
            'Topic',
            'File Name',
            'Created Date'
        ])
        
        # Data
        for i, question in enumerate(self.questions.order_by('bloom_level', 'question_type', 'created_at'), 1):
            writer.writerow([
                i,
                question.question_text,
                question.get_question_type_display(),
                question.get_bloom_level_display(),
                question.get_difficulty_display(),
                question.answer,
                ' | '.join(question.options) if question.options else '',
                question.explanation,
                question.topic or '',
                question.file.filename,
                question.created_at.strftime('%Y-%m-%d %H:%M:%S')
            ])
        
        return output.getvalue()

