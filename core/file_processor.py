# # # core/file_processor.py
# # import fitz  # PyMuPDF
# # from docx import Document
# # import re
# # from typing import Dict, Any
# # from langdetect import detect
# # import os

# # class FileProcessor:
# #     @staticmethod
# #     def extract_text_from_pdf(file_path: str) -> str:
# #         """Extract text from PDF using PyMuPDF"""
# #         text = ""
# #         try:
# #             pdf_document = fitz.open(file_path)
# #             for page_num in range(pdf_document.page_count):
# #                 page = pdf_document[page_num]
# #                 text += page.get_text()
# #             pdf_document.close()
# #         except Exception as e:
# #             raise Exception(f"Error processing PDF: {str(e)}")
# #         return text

# #     @staticmethod
# #     def extract_text_from_docx(file_path: str) -> str:
# #         """Extract text from Word document"""
# #         try:
# #             doc = Document(file_path)
# #             text = ""
# #             for paragraph in doc.paragraphs:
# #                 text += paragraph.text + "\n"
# #         except Exception as e:
# #             raise Exception(f"Error processing DOCX: {str(e)}")
# #         return text

# #     @staticmethod
# #     def extract_text_from_txt(file_path: str) -> str:
# #         """Extract text from plain text file"""
# #         try:
# #             with open(file_path, 'r', encoding='utf-8') as file:
# #                 text = file.read()
# #         except UnicodeDecodeError:
# #             with open(file_path, 'r', encoding='latin-1') as file:
# #                 text = file.read()
# #         except Exception as e:
# #             raise Exception(f"Error processing TXT: {str(e)}")
# #         return text

# #     @staticmethod
# #     def detect_language(text: str) -> str:
# #         """Detect language of the text"""
# #         try:
# #             sample = text[:1000] if len(text) > 1000 else text
# #             return detect(sample)
# #         except:
# #             return "en"

# #     @staticmethod
# #     def detect_math_content(text: str) -> bool:
# #         """Detect if text contains mathematical notation"""
# #         math_patterns = [
# #             r'\$.*?\$',
# #             r'\\frac\{.*?\}\{.*?\}',
# #             r'[∫∑∏√±≤≥≠≈∞]',
# #             r'\b\d+[+\-*/^]\d+\b',
# #         ]
        
# #         for pattern in math_patterns:
# #             if re.search(pattern, text, re.DOTALL):
# #                 return True
# #         return False

# #     @classmethod
# #     def process_file(cls, file_path: str, filename: str) -> Dict[str, Any]:
# #         """Main method to process uploaded file"""
# #         file_ext = filename.lower().split('.')[-1]
        
# #         if file_ext == 'pdf':
# #             text = cls.extract_text_from_pdf(file_path)
# #         elif file_ext in ['docx', 'doc']:
# #             text = cls.extract_text_from_docx(file_path)
# #         elif file_ext in ['txt', 'md']:
# #             text = cls.extract_text_from_txt(file_path)
# #         else:
# #             raise ValueError(f"Unsupported file format: {file_ext}")
        
# #         text = text.strip()
# #         if not text:
# #             raise ValueError("No text content found in file")
        
# #         detected_language = cls.detect_language(text)
# #         math_content_detected = cls.detect_math_content(text)
# #         word_count = len(text.split())
        
# #         return {
# #             "content": text,
# #             "detected_language": detected_language,
# #             "math_content_detected": math_content_detected,
# #             "word_count": word_count
# #         }

# # import fitz  # PyMuPDF
# # from docx import Document
# # import re
# # from typing import Dict, Any, List
# # from langdetect import detect, DetectorFactory
# # import unicodedata
# # import math
# # import sympy as sp

# # # Ensure consistent language detection
# # DetectorFactory.seed = 0

# # class EnhancedFileProcessor:
    
# #     # Language detection patterns
# #     LANGUAGE_PATTERNS = {
# #         'hi': r'[\u0900-\u097F]+',  # Devanagari (Hindi)
# #         'te': r'[\u0C00-\u0C7F]+',  # Telugu
# #         'ur': r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]+',  # Arabic/Urdu
# #         'ar': r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]+',  # Arabic
# #         'bn': r'[\u0980-\u09FF]+',  # Bengali
# #         'ta': r'[\u0B80-\u0BFF]+',  # Tamil
# #         'ml': r'[\u0D00-\u0D7F]+',  # Malayalam
# #         'kn': r'[\u0C80-\u0CFF]+',  # Kannada
# #         'gu': r'[\u0A80-\u0AFF]+',  # Gujarati
# #         'mr': r'[\u0900-\u097F]+',  # Marathi (Devanagari)
# #         'pa': r'[\u0A00-\u0A7F]+',  # Punjabi (Gurmukhi)
# #         'zh': r'[\u4e00-\u9fff\u3400-\u4dbf]+',  # Chinese
# #         'ja': r'[\u3040-\u309f\u30a0-\u30ff\u4e00-\u9fff]+',  # Japanese
# #         'ko': r'[\uac00-\ud7a3\u1100-\u11ff\u3130-\u318f]+',  # Korean
# #         'ru': r'[\u0400-\u04FF]+',  # Cyrillic (Russian)
# #         'th': r'[\u0E00-\u0E7F]+',  # Thai
# #         'vi': r'[àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđĐ]+',  # Vietnamese
# #     }
    
# #     # Mathematical patterns
# #     MATH_PATTERNS = {
# #         'latex_inline': r'\$([^$]+)\$',
# #         'latex_display': r'\$\$([^$]+)\$\$',
# #         'latex_block': r'\\begin\{([^}]+)\}(.*?)\\end\{\1\}',
# #         'equations': r'[a-zA-Z]\s*[=]\s*[^a-zA-Z\s]{2,}',
# #         'fractions': r'\\frac\{([^}]+)\}\{([^}]+)\}',
# #         'symbols': r'[∫∑∏√±≤≥≠≈∞∂∇α-ω]',
# #         'superscript': r'[a-zA-Z0-9]\^[0-9\{\}a-zA-Z]+',
# #         'subscript': r'[a-zA-Z0-9]_[0-9\{\}a-zA-Z]+',
# #         'polynomials': r'[a-zA-Z]\^[0-9]+[\s\+\-][a-zA-Z]',
# #         'functions': r'(sin|cos|tan|log|ln|exp|sqrt)\([^)]+\)',
# #     }
    
# #     @classmethod
# #     def detect_language_advanced(cls, text: str) -> Dict[str, Any]:
# #         """Advanced language detection with script analysis"""
# #         # Basic language detection
# #         try:
# #             primary_lang = detect(text[:2000])
# #         except:
# #             primary_lang = 'en'
        
# #         # Script-based detection
# #         detected_scripts = []
# #         for lang, pattern in cls.LANGUAGE_PATTERNS.items():
# #             if re.search(pattern, text):
# #                 detected_scripts.append(lang)
        
# #         # Determine script type and reading direction
# #         script_type = 'latin'
# #         reading_direction = 'ltr'
        
# #         if any(lang in ['hi', 'mr', 'ne'] for lang in detected_scripts):
# #             script_type = 'devanagari'
# #         elif any(lang in ['te'] for lang in detected_scripts):
# #             script_type = 'telugu'
# #         elif any(lang in ['ur', 'ar'] for lang in detected_scripts):
# #             script_type = 'arabic'
# #             reading_direction = 'rtl'
# #         elif any(lang in ['zh', 'ja', 'ko'] for lang in detected_scripts):
# #             script_type = 'cjk'
# #         elif any(lang in ['bn'] for lang in detected_scripts):
# #             script_type = 'bengali'
# #         elif any(lang in ['ta'] for lang in detected_scripts):
# #             script_type = 'tamil'
        
# #         return {
# #             'primary_language': primary_lang,
# #             'detected_scripts': detected_scripts,
# #             'script_type': script_type,
# #             'reading_direction': reading_direction,
# #             'is_multilingual': len(detected_scripts) > 1
# #         }
    
# #     @classmethod
# #     def detect_mathematical_content(cls, text: str) -> Dict[str, Any]:
# #         """Enhanced mathematical content detection"""
# #         math_info = {
# #             'has_math': False,
# #             'math_expressions': [],
# #             'mathematical_level': 'basic',
# #             'types_detected': []
# #         }
        
# #         for math_type, pattern in cls.MATH_PATTERNS.items():
# #             matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
# #             if matches:
# #                 math_info['has_math'] = True
# #                 math_info['types_detected'].append(math_type)
# #                 math_info['math_expressions'].extend(matches)
        
# #         # Determine mathematical level
# #         if any(t in math_info['types_detected'] for t in ['latex_block', 'fractions']):
# #             math_info['mathematical_level'] = 'advanced'
# #         elif any(t in math_info['types_detected'] for t in ['polynomials', 'functions']):
# #             math_info['mathematical_level'] = 'algebra'
# #         elif any(t in math_info['types_detected'] for t in ['equations', 'symbols']):
# #             math_info['mathematical_level'] = 'intermediate'
        
# #         return math_info
    
# #     @classmethod
# #     def normalize_mathematical_expressions(cls, text: str) -> str:
# #         """Normalize mathematical expressions for processing"""
# #         # Convert common mathematical notation
# #         replacements = {
# #             '×': '*',
# #             '÷': '/',
# #             '²': '^2',
# #             '³': '^3',
# #             '√': 'sqrt',
# #             '∞': 'infinity',
# #             '≤': '<=',
# #             '≥': '>=',
# #             '≠': '!=',
# #             '≈': '~=',
# #             'π': 'pi',
# #             'α': 'alpha',
# #             'β': 'beta',
# #             'γ': 'gamma',
# #             'δ': 'delta',
# #             'θ': 'theta',
# #             'λ': 'lambda',
# #             'μ': 'mu',
# #             'σ': 'sigma',
# #             'φ': 'phi',
# #             'ω': 'omega'
# #         }
        
# #         normalized_text = text
# #         for old, new in replacements.items():
# #             normalized_text = normalized_text.replace(old, new)
        
# #         return normalized_text
    
# #     @classmethod
# #     def extract_mathematical_concepts(cls, text: str) -> List[str]:
# #         """Extract mathematical concepts from text"""
# #         concepts = []
        
# #         # Polynomial detection
# #         if re.search(r'[a-zA-Z]\^[0-9]+', text):
# #             concepts.append('polynomials')
        
# #         # Function detection
# #         functions = re.findall(r'(sin|cos|tan|log|ln|exp|sqrt|abs)\(', text, re.IGNORECASE)
# #         if functions:
# #             concepts.extend(['trigonometry' if f in ['sin', 'cos', 'tan'] else 'functions' for f in functions])
        
# #         # Calculus detection
# #         if re.search(r'(derivative|integral|limit|∫|∂|∇)', text, re.IGNORECASE):
# #             concepts.append('calculus')
        
# #         # Statistics detection
# #         if re.search(r'(mean|median|mode|variance|standard deviation|probability)', text, re.IGNORECASE):
# #             concepts.append('statistics')
        
# #         # Geometry detection
# #         if re.search(r'(triangle|circle|square|rectangle|angle|area|perimeter)', text, re.IGNORECASE):
# #             concepts.append('geometry')
        
# #         return list(set(concepts))
    
# #     @classmethod
# #     def process_file_enhanced(cls, file_path: str, filename: str) -> Dict[str, Any]:
# #         """Enhanced file processing with multi-language and math support"""
# #         # Extract text (using existing methods)
# #         file_ext = filename.lower().split('.')[-1]
        
# #         if file_ext == 'pdf':
# #             text = cls.extract_text_from_pdf(file_path)
# #         elif file_ext in ['docx', 'doc']:
# #             text = cls.extract_text_from_docx(file_path)
# #         elif file_ext in ['txt', 'md']:
# #             text = cls.extract_text_from_txt(file_path)
# #         else:
# #             raise ValueError(f"Unsupported file format: {file_ext}")
        
# #         if not text.strip():
# #             raise ValueError("No text content found in file")
        
# #         # Language analysis
# #         lang_info = cls.detect_language_advanced(text)
        
# #         # Mathematical analysis
# #         math_info = cls.detect_mathematical_content(text)
        
# #         # Content type determination
# #         if math_info['has_math'] and len(math_info['math_expressions']) > 5:
# #             content_type = 'mathematical'
# #         elif math_info['has_math']:
# #             content_type = 'mixed'
# #         else:
# #             content_type = 'text'
        
# #         # Normalize mathematical content if detected
# #         if math_info['has_math']:
# #             text = cls.normalize_mathematical_expressions(text)
        
# #         return {
# #             'content': text,
# #             'detected_language': lang_info['primary_language'],
# #             'script_type': lang_info['script_type'],
# #             'reading_direction': lang_info['reading_direction'],
# #             'is_multilingual': lang_info['is_multilingual'],
# #             'detected_scripts': lang_info['detected_scripts'],
# #             'word_count': len(text.split()),
# #             'math_content_detected': math_info['has_math'],
# #             'mathematical_level': math_info['mathematical_level'],
# #             'math_expressions': math_info['math_expressions'][:10],  # Limit to first 10
# #             'mathematical_concepts': cls.extract_mathematical_concepts(text),
# #             'content_type': content_type,
# #             'content_preview': text[:500] + "..." if len(text) > 500 else text
# #         }
    
# #     # Existing methods (extract_text_from_pdf, etc.) remain the same
# #     @staticmethod
# #     def extract_text_from_pdf(file_path: str) -> str:
# #         text = ""
# #         try:
# #             pdf_document = fitz.open(file_path)
# #             for page_num in range(pdf_document.page_count):
# #                 page = pdf_document[page_num]
# #                 text += page.get_text()
# #             pdf_document.close()
# #         except Exception as e:
# #             raise Exception(f"Error processing PDF: {str(e)}")
# #         return text
    
# #     @staticmethod
# #     def extract_text_from_docx(file_path: str) -> str:
# #         try:
# #             doc = Document(file_path)
# #             text = ""
# #             for paragraph in doc.paragraphs:
# #                 text += paragraph.text + "\n"
# #         except Exception as e:
# #             raise Exception(f"Error processing DOCX: {str(e)}")
# #         return text
    
# #     @staticmethod
# #     def extract_text_from_txt(file_path: str) -> str:
# #         try:
# #             with open(file_path, 'r', encoding='utf-8') as file:
# #                 text = file.read()
# #         except UnicodeDecodeError:
# #             with open(file_path, 'r', encoding='latin-1') as file:
# #                 text = file.read()
# #         except Exception as e:
# #             raise Exception(f"Error processing TXT: {str(e)}")
# #         return text
# # # Add to core/file_processor.py
# # from .language_detector import LanguageDetector

# # class EnhancedFileProcessor(FileProcessor):
    
# #     @classmethod
# #     def process_file_with_language(cls, file_path: str, filename: str) -> Dict[str, Any]:
# #         """Process file with enhanced language detection"""
        
# #         # Extract text using existing methods
# #         file_ext = filename.lower().split('.')[-1]
        
# #         if file_ext == 'pdf':
# #             text = cls.extract_text_from_pdf(file_path)
# #         elif file_ext in ['docx', 'doc']:
# #             text = cls.extract_text_from_docx(file_path)
# #         elif file_ext in ['txt', 'md']:
# #             text = cls.extract_text_from_txt(file_path)
# #         else:
# #             raise ValueError(f"Unsupported file format: {file_ext}")
        
# #         if not text.strip():
# #             raise ValueError("No text content found in file")
        
# #         # Enhanced language detection
# #         lang_info = LanguageDetector.detect_language_advanced(text)
        
# #         # Basic math detection
# #         math_patterns = [r'\$.*?\$', r'[∫∑∏√±≤≥≠≈∞]', r'\b\d+[+\-*/^]\d+\b']
# #         math_content_detected = any(re.search(pattern, text) for pattern in math_patterns)
        
# #         return {
# #             "content": text,
# #             "detected_language": lang_info['language'],
# #             "language_confidence": lang_info['confidence'],
# #             "detected_scripts": lang_info['detected_scripts'],
# #             "script_type": lang_info['script_type'],
# #             "reading_direction": lang_info['reading_direction'],
# #             "font_family": lang_info['font_family'],
# #             "word_count": len(text.split()),
# #             "math_content_detected": math_content_detected,
# #             "content_type": "mathematical" if math_content_detected else "text",
# #             "mathematical_level": "basic" if math_content_detected else "",
# #             "is_supported_language": LanguageDetector.is_supported_language(lang_info['language'])
# #         }






# import os
# import re
# from typing import Dict, Any

# class FileProcessor:
#     @staticmethod
#     def extract_text_from_pdf(file_path: str) -> str:
#         """Extract text from PDF using PyMuPDF"""
#         try:
#             import fitz  # PyMuPDF
#             text = ""
#             pdf_document = fitz.open(file_path)
#             for page_num in range(pdf_document.page_count):
#                 page = pdf_document[page_num]
#                 text += page.get_text()
#             pdf_document.close()
#             return text
#         except ImportError:
#             return "PDF processing requires PyMuPDF. Please install it: pip install PyMuPDF"
#         except Exception as e:
#             raise Exception(f"Error processing PDF: {str(e)}")

#     @staticmethod
#     def extract_text_from_docx(file_path: str) -> str:
#         """Extract text from Word document"""
#         try:
#             from docx import Document
#             doc = Document(file_path)
#             text = ""
#             for paragraph in doc.paragraphs:
#                 text += paragraph.text + "\n"
#             return text
#         except ImportError:
#             return "DOCX processing requires python-docx. Please install it: pip install python-docx"
#         except Exception as e:
#             raise Exception(f"Error processing DOCX: {str(e)}")

#     @staticmethod
#     def extract_text_from_txt(file_path: str) -> str:
#         """Extract text from plain text file"""
#         try:
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 text = file.read()
#             return text
#         except UnicodeDecodeError:
#             try:
#                 with open(file_path, 'r', encoding='latin-1') as file:
#                     text = file.read()
#                 return text
#             except Exception as e:
#                 raise Exception(f"Error processing TXT: {str(e)}")
#         except Exception as e:
#             raise Exception(f"Error processing TXT: {str(e)}")

#     @staticmethod
#     def detect_language(text: str) -> str:
#         """Simple language detection"""
#         try:
#             from langdetect import detect
#             sample = text[:1000] if len(text) > 1000 else text
#             return detect(sample)
#         except ImportError:
#             # Simple fallback language detection
#             if any(char in text for char in ['हिंदी', 'है', 'में', 'की', 'और']):
#                 return 'hi'
#             elif any(char in text for char in ['తెలుగు', 'అని', 'లో', 'కు', 'వా']):
#                 return 'te'
#             elif any(char in text for char in ['اردو', 'کے', 'میں', 'کو', 'ہے']):
#                 return 'ur'
#             elif any(char in text for char in ['العربية', 'في', 'من', 'إلى', 'على']):
#                 return 'ar'
#             else:
#                 return 'en'
#         except:
#             return 'en'

#     @staticmethod
#     def detect_math_content(text: str) -> bool:
#         """Detect if text contains mathematical notation"""
#         math_patterns = [
#             r'\$.*?\$',  # LaTeX math
#             r'\\frac\{.*?\}\{.*?\}',  # Fractions
#             r'[∫∑∏√±≤≥≠≈∞]',  # Math symbols
#             r'\b\d+[+\-*/^]\d+\b',  # Simple expressions
#             r'[xyz]\s*[=]\s*\d+',  # Variables
#         ]
        
#         for pattern in math_patterns:
#             if re.search(pattern, text, re.DOTALL):
#                 return True
#         return False

#     @classmethod
#     def process_file(cls, file_path: str, filename: str) -> Dict[str, Any]:
#         """Main method to process uploaded file"""
#         file_ext = filename.lower().split('.')[-1]
        
#         if file_ext == 'pdf':
#             text = cls.extract_text_from_pdf(file_path)
#         elif file_ext in ['docx', 'doc']:
#             text = cls.extract_text_from_docx(file_path)
#         elif file_ext in ['txt', 'md']:
#             text = cls.extract_text_from_txt(file_path)
#         else:
#             raise ValueError(f"Unsupported file format: {file_ext}")
        
#         text = text.strip()
#         if not text:
#             text = f"Content extracted from {filename}"  # Fallback content
        
#         detected_language = cls.detect_language(text)
#         math_content_detected = cls.detect_math_content(text)
#         word_count = len(text.split())
        
#         # Determine script type and reading direction
#         script_type = 'latin'
#         reading_direction = 'ltr'
#         content_type = 'text'
        
#         if detected_language in ['hi', 'mr', 'ne']:
#             script_type = 'devanagari'
#         elif detected_language == 'te':
#             script_type = 'telugu'
#         elif detected_language in ['ur', 'ar']:
#             script_type = 'arabic'
#             reading_direction = 'rtl'
#         elif detected_language in ['zh', 'ja', 'ko']:
#             script_type = 'chinese'
        
#         if math_content_detected:
#             content_type = 'mathematical'
        
#         return {
#             "content": text,
#             "detected_language": detected_language,
#             "math_content_detected": math_content_detected,
#             "word_count": word_count,
#             "script_type": script_type,
#             "reading_direction": reading_direction,
#             "content_type": content_type,
#             "mathematical_level": "basic" if math_content_detected else ""
#         }

# # Enhanced version with more features
# class EnhancedFileProcessor(FileProcessor):
#     @classmethod
#     def process_file_with_language(cls, file_path: str, filename: str) -> Dict[str, Any]:
#         """Enhanced processing with better language detection"""
#         base_result = cls.process_file(file_path, filename)
        
#         # Additional language confidence and script detection
#         base_result.update({
#             "language_confidence": 0.9,
#             "detected_scripts": [base_result["detected_language"]],
#             "is_supported_language": base_result["detected_language"] in [
#                 'en', 'hi', 'te', 'ur', 'ar', 'es', 'fr', 'de', 'zh', 'ja', 'ko', 'ru'
#             ]
#         })
        
#         return base_result










# core/file_processor.py - Complete Updated File
import os
import re
from typing import Dict, Any

class FileProcessor:
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Enhanced PDF text extraction for multilingual content"""
        try:
            import fitz  # PyMuPDF
            text = ""
            pdf_document = fitz.open(file_path)
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                # Get text with better encoding handling
                page_text = page.get_text("text", flags=fitz.TEXT_PRESERVE_WHITESPACE)
                
                # If no text found, try different extraction method
                if not page_text.strip():
                    # Try extracting text blocks
                    text_blocks = page.get_text("blocks")
                    page_text = ""
                    for block in text_blocks:
                        if len(block) >= 5:  # Text block
                            page_text += block[4] + " "
                
                # Clean and add page text
                if page_text.strip():
                    text += page_text.strip() + "\n\n"
            
            pdf_document.close()
            
            # Clean up the extracted text
            text = text.strip()
            if not text:
                raise ValueError("No text could be extracted from PDF")
            
            # Remove excessive whitespace while preserving structure
            text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Multiple newlines to double
            text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
            
            return text
            
        except ImportError:
            return "PDF processing requires PyMuPDF. Please install it: pip install PyMuPDF"
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Enhanced Word document text extraction"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text.strip() + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text.strip() + " "
                    text += "\n"
            
            return text.strip()
            
        except ImportError:
            return "DOCX processing requires python-docx. Please install it: pip install python-docx"
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")

    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Enhanced plain text file extraction"""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return text.strip()
        except UnicodeDecodeError:
            try:
                # Try UTF-8 with error handling
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    text = file.read()
                return text.strip()
            except Exception:
                try:
                    # Try latin-1 as fallback
                    with open(file_path, 'r', encoding='latin-1') as file:
                        text = file.read()
                    return text.strip()
                except Exception as e:
                    raise Exception(f"Error processing TXT: {str(e)}")
        except Exception as e:
            raise Exception(f"Error processing TXT: {str(e)}")

    @staticmethod
    def detect_language(text: str) -> str:
        """Enhanced language detection with script analysis"""
        if not text or len(text.strip()) < 3:
            return 'en'
        
        # Script-based detection patterns
        script_patterns = {
            'hi': r'[\u0900-\u097F]+',  # Devanagari (Hindi)
            'te': r'[\u0C00-\u0C7F]+',  # Telugu
            'ur': r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]+',  # Arabic/Urdu
            'ar': r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]+',  # Arabic
            'bn': r'[\u0980-\u09FF]+',  # Bengali
            'ta': r'[\u0B80-\u0BFF]+',  # Tamil
            'ml': r'[\u0D00-\u0D7F]+',  # Malayalam
            'kn': r'[\u0C80-\u0CFF]+',  # Kannada
            'gu': r'[\u0A80-\u0AFF]+',  # Gujarati
            'mr': r'[\u0900-\u097F]+',  # Marathi (Devanagari)
            'pa': r'[\u0A00-\u0A7F]+',  # Punjabi (Gurmukhi)
        }
        
        # Check for script patterns
        script_counts = {}
        for lang_code, pattern in script_patterns.items():
            matches = re.findall(pattern, text)
            if matches:
                script_counts[lang_code] = len(''.join(matches))
        
        # If script patterns found, use the most frequent one
        if script_counts:
            detected_lang = max(script_counts, key=script_counts.get)
            print(f"Language detected via script analysis: {detected_lang}")
            return detected_lang
        
        # Fallback to langdetect if available
        try:
            from langdetect import detect
            sample = text[:2000] if len(text) > 2000 else text
            detected = detect(sample)
            print(f"Language detected via langdetect: {detected}")
            return detected
        except ImportError:
            print("langdetect not available, checking for common words...")
            # Simple word-based detection as last resort
            text_lower = text.lower()
            
            # Hindi words
            if any(word in text for word in ['हिंदी', 'है', 'में', 'की', 'और', 'का', 'के', 'से', 'को', 'पर']):
                return 'hi'
            # Telugu words
            elif any(word in text for word in ['తెలుగు', 'అని', 'లో', 'కు', 'వా', 'ని', 'గా', 'కి', 'చే']):
                return 'te'
            # Urdu words
            elif any(word in text for word in ['اردو', 'کے', 'میں', 'کو', 'ہے', 'اور', 'کا', 'سے', 'پر']):
                return 'ur'
            # Arabic words
            elif any(word in text for word in ['العربية', 'في', 'من', 'إلى', 'على', 'هذا', 'التي', 'كان']):
                return 'ar'
            else:
                return 'en'
        except Exception:
            return 'en'

    @staticmethod
    def detect_math_content(text: str) -> bool:
        """Enhanced mathematical content detection"""
        if not text:
            return False
        
        math_patterns = [
            r'\$.*?\$'     # LaTeX math
            r'\\frac\{.*?\}\{.*?\}',  # Fractions
            r'[∫∑∏√±≤≥≠≈∞∂∇α-ω]',  # Math symbols
            r'\b\d+[+\-*/^]\d+\b',  # Simple expressions
            r'[xyz]\s*[=]\s*\d+',  # Variables
            r'\\[a-zA-Z]+\{.*?\}',  # LaTeX commands
            r'\b(equation|formula|theorem|proof|calculate)\b',  # Math keywords
        ]
        
        for pattern in math_patterns:
            if re.search(pattern, text, re.DOTALL | re.IGNORECASE):
                return True
        return False

    @classmethod
    def process_file(cls, file_path: str, filename: str) -> Dict[str, Any]:
        """Main method to process uploaded file with enhanced language support"""
        if not os.path.exists(file_path):
            raise ValueError(f"File not found: {file_path}")
        
        file_ext = filename.lower().split('.')[-1]
        
        if file_ext == 'pdf':
            text = cls.extract_text_from_pdf(file_path)
        elif file_ext in ['docx', 'doc']:
            text = cls.extract_text_from_docx(file_path)
        elif file_ext in ['txt', 'md']:
            text = cls.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        text = text.strip()
        if not text:
            raise ValueError("No text content found in file")
        
        # Enhanced language detection
        detected_language = cls.detect_language(text)
        math_content_detected = cls.detect_math_content(text)
        word_count = len(text.split())
        
        # Determine script type and reading direction
        script_type = 'latin'
        reading_direction = 'ltr'
        content_type = 'text'
        
        if detected_language in ['hi', 'mr', 'ne']:
            script_type = 'devanagari'
        elif detected_language == 'te':
            script_type = 'telugu'
        elif detected_language in ['ur', 'ar']:
            script_type = 'arabic'
            reading_direction = 'rtl'
        elif detected_language in ['zh', 'ja', 'ko']:
            script_type = 'chinese'
        
        if math_content_detected:
            content_type = 'mathematical'
        
        print(f"File processed: {filename}")
        print(f"Language detected: {detected_language}")
        print(f"Word count: {word_count}")
        print(f"Math content: {math_content_detected}")
        print(f"Content preview: {text[:200]}...")
        
        return {
            "content": text,
            "detected_language": detected_language,
            "math_content_detected": math_content_detected,
            "word_count": word_count,
            "script_type": script_type,
            "reading_direction": reading_direction,
            "content_type": content_type,
            "mathematical_level": "basic" if math_content_detected else ""
        }

# Enhanced version with more features
class EnhancedFileProcessor(FileProcessor):
    @classmethod
    def process_file_with_language(cls, file_path: str, filename: str) -> Dict[str, Any]:
        """Enhanced processing with better language detection"""
        base_result = cls.process_file(file_path, filename)
        
        # Additional enhancements
        text = base_result["content"]
        detected_language = base_result["detected_language"]
        
        # Language confidence estimation
        confidence = 0.9 if base_result["script_type"] != 'latin' else 0.7
        
        # Enhanced script detection
        detected_scripts = [detected_language]
        if base_result["script_type"] != 'latin':
            detected_scripts.append(base_result["script_type"])
        
        # Check if language is supported
        supported_languages = [
            'en', 'hi', 'te', 'ur', 'ar', 'es', 'fr', 'de', 'zh', 'ja', 'ko', 'ru',
            'pt', 'it', 'bn', 'ta', 'ml', 'kn', 'gu', 'mr', 'pa', 'or', 'as'
        ]
        is_supported = detected_language in supported_languages
        
        # Content analysis
        content_analysis = cls._analyze_content_structure(text, detected_language)
        
        base_result.update({
            "language_confidence": confidence,
            "detected_scripts": list(set(detected_scripts)),
            "is_supported_language": is_supported,
            "content_analysis": content_analysis
        })
        
        return base_result
    
    @classmethod
    def _analyze_content_structure(cls, text: str, language: str) -> Dict[str, Any]:
        """Analyze content structure for better question generation"""
        lines = text.split('\n')
        sentences = re.split(r'[।\.\!\?]', text)
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        # Extract potential topics/entities
        if language == 'hi':
            # Extract Hindi words that might be names or concepts
            entities = re.findall(r'[\u0900-\u097F]{3,}', text)
        elif language == 'te':
            entities = re.findall(r'[\u0C00-\u0C7F]{3,}', text)
        elif language == 'ur':
            entities = re.findall(r'[\u0600-\u06FF]{3,}', text)
        else:
            # For English and other Latin scripts
            entities = re.findall(r'\b[A-Z][a-z]{2,}\b', text)
        
        # Find dates and numbers
        dates = re.findall(r'\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}|\d{4}', text)
        numbers = re.findall(r'\b\d+\b', text)
        
        return {
            'line_count': len(lines),
            'sentence_count': len([s for s in sentences if s.strip()]),
            'paragraph_count': len(paragraphs),
            'entities': list(set(entities))[:10],  # Top 10 unique entities
            'dates': list(set(dates))[:5],  # Top 5 unique dates
            'numbers': list(set(numbers))[:5],  # Top 5 unique numbers
            'avg_sentence_length': sum(len(s.split()) for s in sentences if s.strip()) / max(len(sentences), 1),
            'longest_paragraph': max(paragraphs, key=len) if paragraphs else "",
        }